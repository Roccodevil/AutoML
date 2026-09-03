import pandas as pd
import io
import os
import json
import math
import time
import pdfplumber
import docx
import re
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from core.llm_services import llm_powerful_api
from huggingface_hub import HfApi
from datasets import load_dataset
from langchain_community.tools.tavily_search import TavilySearchResults

# --- SAFE IMPORT: Kaggle ---
# --- SEARCH & DATASETS ---
try:
    from huggingface_hub import HfApi
    from datasets import load_dataset
except ImportError:
    HfApi = None

try: 
    from kaggle.api.kaggle_api_extended import KaggleApi
except ImportError: 
    KaggleApi = None

class DataAgent:
    def __init__(self):
        # 1. SPECIALIZED PROMPT: Handles both Tables and Vertical Key-Value lists
        self.kv_extraction_prompt = ChatPromptTemplate.from_template(
            """You are an Expert Data Extraction AI.
            Your task is to convert unstructured text chunks into a structured JSON list.

            ANALYSIS STEPS:
            1. Scan the text for repeating record patterns.
            2. Detect the format: Is it a Table (row-by-row) or Vertical List (Key: Value blocks)?
            3. Identify common fields (e.g., "Id", "SepalLength", "Species", "Price").
            4. Extract EVERY record found in this chunk.
            
            CRITICAL OUTPUT RULES:
            - Return a **JSON LIST** of objects. `[ {{...}}, {{...}} ]`
            - Do NOT nest the list.
            - If a record is clearly incomplete (cut off), skip it.
            - CLEANUP: Remove currency symbols ($, €) or commas in numbers (e.g., "1,000" -> 1000).
            
            EXAMPLE INPUT (Vertical Style):
            "Id: 1
             SepalLength: 5.1
             Species: Setosa
             Id: 2
             SepalLength: 4.9
             Species: Setosa"

            EXAMPLE OUTPUT:
            [
                {{ "Id": 1, "SepalLength": 5.1, "Species": "Setosa" }},
                {{ "Id": 2, "SepalLength": 4.9, "Species": "Setosa" }}
            ]

            Raw Text Chunk:
            {text_data}
            """
        )
        self.kv_chain = self.kv_extraction_prompt | llm_powerful_api | JsonOutputParser()

        # 2. GENERATION PROMPT (Updated with Live Search Context)
        self.gen_search_prompt = ChatPromptTemplate.from_template(
            """You are an Expert Data Generator.
            The user wants a dataset about: "{description}"
            
            REAL-WORLD CONTEXT (Recently Searched):
            {context}
            
            TASK:
            Generate a realistic, structured dataset (JSON list) based on the description and the provided context.
            - Use REAL names, prices, dates, or specs found in the context.
            - If context is thin, infer realistic values based on the data profile.
            - Generate at least 20-30 rows to allow for analysis.
            
            Return ONLY valid JSON: `[ {{"Column": Value, ...}}, ... ]`
            """
        )
        self.gen_chain = self.gen_search_prompt | llm_powerful_api | JsonOutputParser()

        self.hf_api = HfApi()
        
        # Authenticate Kaggle if available
        self.kaggle_api = None
        k_user = os.environ.get('KAGGLE_USERNAME')
        k_key = os.environ.get('KAGGLE_KEY')
        local_creds = os.path.expanduser("~/.kaggle/kaggle.json")
        if KaggleApi and ((k_user and k_key) or os.path.exists(local_creds)):
            try:
                self.kaggle_api = KaggleApi()
                self.kaggle_api.authenticate()
            except: pass

        try: self.web_search = TavilySearchResults(max_results=3)
        except: self.web_search = None

    def _load_structured_file(self, filepath, ext):
        """Loads standard formats like Excel, XML, JSON."""
        try:
            if ext in ['.xlsx', '.xls']: return pd.read_excel(filepath)
            elif ext == '.json':
                try: return pd.read_json(filepath)
                except: 
                    with open(filepath) as f: return pd.json_normalize(json.load(f))
            elif ext == '.xml': return pd.read_xml(filepath)
            elif ext == '.parquet': return pd.read_parquet(filepath)
        except Exception as e:
            print(f"   [Error] Structured load failed: {str(e)}")
        return None

    def _fallback_regex_parsing(self, text_chunk):
        """
        FAILOVER: If LLM dies (Rate Limit), this regex extracts 'Key: Value' patterns.
        It detects new records when keys repeat (e.g., seeing 'Id' twice).
        """
        print("   [Agent 1] Triggering REGEX FALLBACK (LLM Rate Limit Bypass)...")
        
        lines = text_chunk.split('\n')
        records = []
        current_record = {}
        
        for line in lines:
            # Check for Key: Value pattern
            match = re.match(r'^\s*([a-zA-Z0-9_]+)\s*:\s*(.+)$', line)
            if match:
                key, val = match.groups()
                key = key.strip()
                val = val.strip()
                
                # If key already exists in current_record, assume new record started
                if key in current_record:
                    records.append(current_record)
                    current_record = {}
                
                current_record[key] = val
            elif not line.strip():
                # Empty line can also signal end of record if we have data
                if current_record:
                    records.append(current_record)
                    current_record = {}
        
        # Append the last one
        if current_record:
            records.append(current_record)
            
        return pd.DataFrame(records)

    def _extract_from_doc(self, filepath, ext):
        """Robust extraction for PDFs, DOCX, and TXT files."""
        print(f"   [Agent 1] Extracting unstructured data from {ext}...")
        text_content = ""
        tables = []

        # 1. Read Content
        if ext == '.pdf':
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_table()
                    if extracted: tables.extend(extracted)
                    text_content += (page.extract_text() or "") + "\n"
        elif ext == '.docx':
            doc = docx.Document(filepath)
            for table in doc.tables:
                data = [[cell.text for cell in row.cells] for row in table.rows]
                tables.extend(data)
            text_content += "\n".join([p.text for p in doc.paragraphs])
        elif ext == '.txt':
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()

        # 2. Heuristic: If we found many tables, use them directly (faster/safer)
        if len(tables) > 5: 
            print("   [Agent 1] Structure detected in tables. Skipping LLM extraction.")
            headers = tables[0]
            # Fix duplicate headers
            headers = [f"{h}_{i}" if headers.count(h) > 1 else h for i, h in enumerate(headers)]
            return pd.DataFrame(tables[1:], columns=headers)

        # 3. LLM Parsing Loop with FALLBACK
        CHUNK_SIZE = 3000
        OVERLAP = 300 
        total_length = len(text_content)
        extracted_dfs = []
        current_pos = 0
        chunk_counter = 0
        
        while current_pos < total_length:
            chunk_counter += 1
            end_pos = min(current_pos + CHUNK_SIZE, total_length)
            chunk = text_content[current_pos:end_pos]
            
            # Retry mechanism
            success = False
            retries = 1 # Keep low to fail fast to regex if needed
            
            for attempt in range(retries + 1):
                try:
                    if chunk_counter > 1: time.sleep(0.5) # Rate limit protection
                    
                    # Try LLM
                    json_data = self.kv_chain.invoke({"text_data": chunk})
                    
                    chunk_df = None
                    if isinstance(json_data, list):
                        chunk_df = pd.DataFrame(json_data)
                    elif isinstance(json_data, dict):
                        # Handle case where LLM returns a dict wrapping the list
                        if any(isinstance(v, list) for v in json_data.values()):
                            for k, v in json_data.items():
                                if isinstance(v, list):
                                    chunk_df = pd.DataFrame(v)
                                    break
                        else:
                            chunk_df = pd.DataFrame([json_data])
                    
                    if chunk_df is not None and not chunk_df.empty:
                        extracted_dfs.append(chunk_df)
                        success = True
                        break # Success, exit retry loop
                except Exception as e:
                    err_msg = str(e)
                    # Detect Rate Limit (429) specifically
                    if "429" in err_msg or "Rate limit" in err_msg:
                        print(f"   [Agent 1] Rate Limit Hit on Chunk {chunk_counter}. Switching to Fallback.")
                        break # Stop retrying LLM, go to fallback
                    time.sleep(1)

            # If LLM failed (Rate Limit or Parsing Error), use Regex Fallback
            if not success:
                fallback_df = self._fallback_regex_parsing(chunk)
                if fallback_df is not None and not fallback_df.empty:
                    extracted_dfs.append(fallback_df)
                    print(f"   [Agent 1] Recovered Chunk {chunk_counter} using Regex.")
                else:
                    print(f"   [Warning] Chunk {chunk_counter} dropped (No LLM, No Regex match).")

            current_pos += (CHUNK_SIZE - OVERLAP)

        if len(extracted_dfs) > 0:
            full_df = pd.concat(extracted_dfs, ignore_index=True)
            full_df = full_df.drop_duplicates()
            print(f"   [Agent 1] Extraction complete. Rows retrieved: {len(full_df)}")
            return full_df
        
        raise ValueError("Could not extract any valid data from the document.")

    def search_online(self, query, max_results=5):
        results = []
        if self.kaggle_api:
            try:
                k_res = self.kaggle_api.dataset_list(search=query, sort_by='votes', file_type='csv')
                for d in k_res[:max_results]:
                    results.append({'source': 'kaggle', 'id': d.ref, 'display': f"Kaggle: {d.title}"})
            except: pass
        try:
            hf_res = list(self.hf_api.list_datasets(search=query, full=True, limit=50))
            for info in sorted(hf_res, key=lambda x: x.downloads or 0, reverse=True)[:max_results]:
                 if info.downloads > 10: results.append({'source': 'huggingface', 'id': info.id, 'display': f"HF: {info.id}"})
        except: pass
        return results

    def download_selected(self, source, dataset_id, download_path):
        os.makedirs(download_path, exist_ok=True)
        if source == 'kaggle':
            self.kaggle_api.dataset_download_files(dataset_id, path=download_path, unzip=True)
            csvs = [f for f in os.listdir(download_path) if f.endswith('.csv')]
            if not csvs: return None
            return pd.read_csv(os.path.join(download_path, max(csvs, key=lambda f: os.path.getsize(os.path.join(download_path, f)))))
        elif source == 'huggingface':
            ds = load_dataset(dataset_id)
            return ds[list(ds.keys())[0]].to_pandas()
        return None

    def _smart_numeric_conversion(self, df):
        """
        Smartly converts columns to numeric. 
        If a column is mostly numbers (e.g. >70%), it enforces numeric 
        conversion. This fixes the 'Object' type issue for numeric targets.
        """
        for col in df.columns:
            # Attempt to convert to numeric
            converted = pd.to_numeric(df[col], errors='coerce')
            
            # Calculate how many values became NaN that weren't NaN before
            non_na_count_before = df[col].notna().sum()
            non_na_count_after = converted.notna().sum()
            
            if non_na_count_before == 0: continue # Skip empty columns

            # Retention Rate: If we keep >70% of data, it's likely a number column
            retention_rate = non_na_count_after / non_na_count_before
            
            if retention_rate > 0.70:
                df[col] = converted
            else:
                # Keep as object/string if too much data would be lost
                try:
                    df[col] = df[col].astype(str).str.strip()
                except: pass
        return df

    def run(self, state):
        print("-> Agent 1: Data Acquisition...")
        mode = state.get('acquisition_mode', 'upload')
        inp = state.get('acquisition_input')
        df = None

        # 1. LOAD DATA
        if mode == "upload":
            if not os.path.exists(inp): raise FileNotFoundError(f"File missing: {inp}")
            ext = os.path.splitext(inp)[1].lower()
            if ext == '.csv': df = pd.read_csv(inp)
            elif ext in ['.xlsx', '.xls', '.json', '.xml', '.parquet']: df = self._load_structured_file(inp, ext)
            elif ext in ['.pdf', '.docx', '.txt']: df = self._extract_from_doc(inp, ext)
            else: raise ValueError(f"Unsupported format: {ext}")

# --- MODE 2: SEARCH (HuggingFace / Kaggle) ---
        elif mode == 'search':
            query = inp
            print(f"   [Agent 1] Online Search for: '{query}'")
            
            # A. Hugging Face Search
            if df is None and HfApi:
                try:
                    print("   ...Querying Hugging Face Hub...")
                    api = HfApi()
                    # Search for datasets, prioritize CSVs
                    datasets = list(api.list_datasets(search=query, limit=5, filter="csv"))
                    
                    if datasets:
                        target_ds = datasets[0].id
                        print(f"   [Agent 1] Found HF Dataset: {target_ds}")
                        # Load the first split (usually 'train') and convert to Pandas
                        ds_data = load_dataset(target_ds, split='train')
                        df = ds_data.to_pandas()
                    else:
                        print("   [Agent 1] No CSV datasets found on Hugging Face.")
                except Exception as e: 
                    print(f"   [Agent 1] HF Search Error: {e}")

            # B. Kaggle Search (Fallback)
            if df is None and KaggleApi:
                try:
                    print("   ...Querying Kaggle...")
                    api = KaggleApi()
                    api.authenticate() # Requires ~/.kaggle/kaggle.json or ENV vars
                    
                    # Search and get the first CSV dataset
                    datasets = api.dataset_list(search=query, file_type='csv')
                    
                    if datasets:
                        target = datasets[0].ref
                        print(f"   [Agent 1] Downloading Kaggle Dataset: {target}")
                        
                        # Download to a temp folder
                        out_dir = os.path.join(os.getcwd(), "data", "kaggle_temp")
                        os.makedirs(out_dir, exist_ok=True)
                        api.dataset_download_files(target, path=out_dir, unzip=True)
                        
                        # Find and load the first CSV file
                        for root, dirs, files in os.walk(out_dir):
                            for f in files:
                                if f.endswith(".csv"):
                                    df = pd.read_csv(os.path.join(root, f))
                                    break
                            if df is not None: break
                except Exception as e: 
                    print(f"   [Agent 1] Kaggle Error: {e}")

            if df is None:
                raise ValueError("Online search returned no usable data from HF or Kaggle.")
        elif mode == "generate":
            print(f"   [Agent 1] Generating data with live search context...")
            search_context = "No online info available."
            
            # --- USE TAVILY FOR LIVE CONTEXT ---
            if self.web_search:
                try:
                    raw_results = self.web_search.invoke(inp)
                    # Aggregate snippets
                    search_context = "\n".join([f"- {r['content']}" for r in raw_results])
                    print(f"   [Agent 1] Retrieved context length: {len(search_context)}")
                except Exception as e:
                    print(f"   [Warning] Web search failed: {e}")

            # Generate Data using Context
            json_data = self.gen_chain.invoke({
                "description": inp, 
                "context": search_context[:3000] # Limit context size
            })
            df = pd.DataFrame(json_data)

        elif mode == "download_selected":
             dl_dir = os.path.join(state.get('results_dir', '.'), "downloaded_data")
             df = self.download_selected(inp['source'], inp['id'], dl_dir)

        if df is not None:
            # 2. SANITIZE
            original_len = len(df)
            df.dropna(how='all', inplace=True) 
            df.dropna(axis=1, how='all', inplace=True) 
            
            if df.empty:
                raise ValueError("Agent 1 Error: Dataset is empty after loading.")

            # 3. HEAL HEADERS (If raw data lacks headers)
            # Only promote first row to header if headers look like default index (0, 1, 2...)
            if len(df) > 1 and all(str(c).isdigit() for c in df.columns):
                first_row = df.iloc[0]
                # If first row contains strings and second row is numbers, it's likely a header
                if any(isinstance(x, str) for x in first_row):
                    print("   [Agent 1] Detected header in first row. Promoting...")
                    df.columns = first_row
                    df = df[1:].reset_index(drop=True)
            
            # 4. FIX TYPES (Crucial for AutoML)
            print("   [Agent 1] Auto-converting numeric columns...")
            df = self._smart_numeric_conversion(df)

            # 5. SHUFFLE (Crucial for Time-Series avoidance in simple split)
            df = df.sample(frac=1, random_state=42).reset_index(drop=True)

            print(f"   Data Loaded. Shape: {df.shape}")
            
            # --- PROFESSIONAL STATE UPDATE ---
            state['raw_df'] = df.copy() # Keep backup
            state['current_data'] = df  # Main Unified Stream
            
            try:
                state['data_preview_html'] = df.head(50).to_html(classes='table', border=0, index=False)
            except: 
                state['data_preview_html'] = "<p>Preview unavailable</p>"
                
            state['data_shape'] = str(df.shape)
            
            # --- STANDALONE OUTPUT FIX ---
            # Ensure the UI shows a success message if the pipeline is stopped here
            state['final_message'] = (
                f"✅ Data Acquisition Complete.\n"
                f"Shape: {df.shape}\n"
                f"Columns: {', '.join(df.columns[:5])}..."
            )
            # -----------------------------
            
        return state